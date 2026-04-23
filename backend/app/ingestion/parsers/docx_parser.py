import io
from app.ingestion.chunker import RawChunk, chunk_markdown


def parse_docx(file_bytes: bytes, doc_title: str) -> list[RawChunk]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    # Paragraphs — convert Word heading styles to markdown headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            level = min(max(level, 1), 6)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    # Tables — render as markdown pipe tables
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header_cells = [c.text.strip() for c in rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    content = "\n\n".join(lines)
    return chunk_markdown(content, source_title=doc_title, source_path=doc_title)
